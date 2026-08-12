"""
contract_builder.py — legalkit-vn v2.0
Render hợp đồng từ JSON template + user input → xuất file .docx

Chuẩn trình bày: Nghị định 30/2020/NĐ-CP, Phụ lục I
  - Font: Times New Roman, Unicode
  - Lề: Trái 3.0cm, Phải 2.0cm, Trên 2.0cm, Dưới 2.0cm
  - Quốc hiệu/Tiêu ngữ đúng chuẩn
  - Khối ký 2 cột cân đối
  - Số trang header (từ trang 2)

Cách dùng:
    python contract_builder.py --template hop-dong-lao-dong.json --input data.json --output output.docx
    python contract_builder.py --template hop-dong-lao-dong.json   (interactive mode)
"""

import json
import sys
import os
import re
import argparse
from datetime import datetime
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

SKILL_DIR = Path(__file__).parent.parent
CONTRACTS_DIR = SKILL_DIR / "templates" / "contracts"
OUTPUT_DIR = SKILL_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


# ─────────────────────────────────────────────
# SECTION 1: Load & Validate
# ─────────────────────────────────────────────

def load_template(template_name: str) -> dict:
    if not template_name.endswith(".json"):
        template_name += ".json"
    path = CONTRACTS_DIR / template_name
    if not path.exists():
        available = [f.name for f in CONTRACTS_DIR.glob("*.json")]
        raise FileNotFoundError(
            f"Template '{template_name}' khong tim thay.\n"
            f"Co san: {', '.join(available)}"
        )
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def collect_inputs_interactive(template: dict) -> dict:
    print(f"\n=== {template['name']} ===")
    values = {}
    current_section = None
    for field in template.get("fields", []):
        condition = field.get("condition")
        if condition and values.get(condition["key"]) != condition["equals"]:
            continue
        section = field.get("section", "")
        if section != current_section:
            print(f"\n--- {section} ---")
            current_section = section
        tier = field.get("tier", "optional")
        label = field["label"]
        tier_badge = {"required": "[BAT BUOC]", "default": "[Co mac dinh luat]", "optional": "[Tuy chon]"}.get(tier, "")
        print(f"\n{tier_badge} {label}")
        if field.get("help"):
            print(f"  Goi y: {field['help']}")
        if field.get("placeholder"):
            print(f"  Vi du: {field['placeholder']}")
        default_text = field.get("defaultText", "")
        if tier == "default" and default_text:
            print(f"  Mac dinh: {default_text[:80]}...")
        while True:
            user_input = input("  Nhap: ").strip()
            if not user_input:
                if tier == "required":
                    print("  Truong nay bat buoc!")
                    continue
                elif tier == "default":
                    values[field["key"]] = default_text
                else:
                    values[field["key"]] = ""
            else:
                values[field["key"]] = user_input
            break
    return values


def run_lint(template: dict, values: dict) -> list:
    warnings = []
    for rule in template.get("lint", []):
        field_key = rule.get("field")
        op = rule.get("op")
        threshold = rule.get("value")
        message = rule.get("message", "Canh bao phap ly")
        legal_ref = rule.get("legalRef", "")
        severity = rule.get("severity", "warning")
        val = values.get(field_key, "")
        if not val:
            continue
        try:
            numeric_val = float(str(val).replace(",", "").replace("%", "").strip())
        except (ValueError, TypeError):
            continue
        triggered = False
        if op == "max" and numeric_val > threshold:
            triggered = True
        elif op == "min" and numeric_val < threshold:
            triggered = True
        elif op == "note":
            triggered = True
        if triggered:
            warnings.append({"severity": severity, "field": field_key,
                             "message": message, "legalRef": legal_ref})
    return warnings


def render_text(template_str: str, values: dict) -> str:
    def replacer(match):
        key = match.group(1)
        return values.get(key, f"[THIEU: {key}]")
    return re.sub(r"\{\{(\w+)\}\}", replacer, template_str)


# ─────────────────────────────────────────────
# SECTION 2: Document Renderer (NĐ 30/2020)
# ─────────────────────────────────────────────

def _add_run(para, text: str, bold=False, italic=False, size_pt=13, font="Times New Roman", underline=False, color=None):
    """Helper: thêm run vào paragraph với format chuẩn NĐ 30."""
    from docx.shared import Pt, RGBColor
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Buộc font Unicode cho tiếng Việt
    from docx.oxml.ns import qn
    rFonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), font)
    return run


def _set_para_format(para, alignment, space_before_pt=0, space_after_pt=6, line_spacing_pt=17, first_indent_cm=0):
    """Helper: định dạng paragraph chuẩn NĐ 30."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_LINE_SPACING
    para.alignment = alignment
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line_spacing_pt)
    if first_indent_cm > 0:
        fmt.first_line_indent = Cm(first_indent_cm)


def build_docx_nd30(template: dict, values: dict) -> "Document":
    """
    Render hợp đồng ra Document object theo chuẩn NĐ 30/2020/NĐ-CP.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # ── 2.1 PAGE SETUP (NĐ 30: Lề trái 3cm, phải 2cm, trên 2cm, dưới 2cm) ──
    for sect in doc.sections:
        sect.page_width  = Cm(21.0)
        sect.page_height = Cm(29.7)
        sect.top_margin    = Cm(2.0)
        sect.bottom_margin = Cm(2.0)
        sect.left_margin   = Cm(3.0)
        sect.right_margin  = Cm(2.0)

    # ── 2.2 HEADER: số trang từ trang 2 ──
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    _set_para_format(hp, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 14)
    # Dùng field PAGE
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r_elem = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "26")
    rpr.append(sz)
    r_elem.append(rpr)
    fld.append(r_elem)
    hp._p.append(fld)

    # ── 2.3 HEADER TABLE (NĐ 30: 2 columns x 2 rows, hidden borders) ──
    header_table = doc.add_table(rows=2, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.style = "Table Grid"

    # Remove borders
    tbl = header_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_elem = OxmlElement(f"w:{side}")
        border_elem.set(qn("w:val"), "none")
        border_elem.set(qn("w:sz"), "0")
        border_elem.set(qn("w:space"), "0")
        border_elem.set(qn("w:color"), "auto")
        tblBorders.append(border_elem)
    tblPr.append(tblBorders)

    # Set column widths (3500 dxa and 5571 dxa)
    col_left_width = Cm(6.17)
    col_right_width = Cm(9.83)
    for row in header_table.rows:
        row.cells[0].width = col_left_width
        row.cells[1].width = col_right_width

    # Row 0, Cell 0: Governing Agency (if any) + Issuing Agency + 1/3 Underline
    cell_d1_l = header_table.rows[0].cells[0]
    cell_d1_l.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for p in cell_d1_l.paragraphs:
        p._element.getparent().remove(p._element)

    ten_nsdld = values.get("ten_nsdld", "")
    p_gov = cell_d1_l.add_paragraph()
    _set_para_format(p_gov, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 17)
    _add_run(p_gov, ten_nsdld.upper(), bold=True, size_pt=13)

    # Underline 1/3
    p_line_l = cell_d1_l.add_paragraph()
    _set_para_format(p_line_l, WD_ALIGN_PARAGRAPH.CENTER, 2, 0, 12)
    pPr_l = p_line_l._p.get_or_add_pPr()
    pBdr_l = OxmlElement("w:pBdr")
    top_l = OxmlElement("w:top")
    top_l.set(qn("w:val"), "single")
    top_l.set(qn("w:sz"), "6")
    top_l.set(qn("w:space"), "1")
    top_l.set(qn("w:color"), "000000")
    pBdr_l.append(top_l)
    pPr_l.append(pBdr_l)
    p_line_l.paragraph_format.left_indent = Cm(2.0)
    p_line_l.paragraph_format.right_indent = Cm(2.0)

    # Row 0, Cell 1: Quốc hiệu + Tiêu ngữ + Underline
    cell_d1_r = header_table.rows[0].cells[1]
    cell_d1_r.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for p in cell_d1_r.paragraphs:
        p._element.getparent().remove(p._element)

    p_qh = cell_d1_r.add_paragraph()
    _set_para_format(p_qh, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 17)
    _add_run(p_qh, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size_pt=13)

    p_tn = cell_d1_r.add_paragraph()
    _set_para_format(p_tn, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 17)
    _add_run(p_tn, "Độc lập - Tự do - Hạnh phúc", bold=True, size_pt=14)

    # Underline matching length
    p_line_r = cell_d1_r.add_paragraph()
    _set_para_format(p_line_r, WD_ALIGN_PARAGRAPH.CENTER, 2, 0, 12)
    pPr_r = p_line_r._p.get_or_add_pPr()
    pBdr_r = OxmlElement("w:pBdr")
    top_r = OxmlElement("w:top")
    top_r.set(qn("w:val"), "single")
    top_r.set(qn("w:sz"), "6")
    top_r.set(qn("w:space"), "1")
    top_r.set(qn("w:color"), "000000")
    pBdr_r.append(top_r)
    pPr_r.append(pBdr_r)
    p_line_r.paragraph_format.left_indent = Cm(1.94)
    p_line_r.paragraph_format.right_indent = Cm(1.94)

    # Row 1, Cell 0: Số ký hiệu
    cell_d2_l = header_table.rows[1].cells[0]
    cell_d2_l.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for p in cell_d2_l.paragraphs:
        p._element.getparent().remove(p._element)

    p_so = cell_d2_l.add_paragraph()
    _set_para_format(p_so, WD_ALIGN_PARAGRAPH.CENTER, 6, 0, 17)
    _add_run(p_so, "Số:         /HĐLĐ", size_pt=13)

    # Row 1, Cell 1: Địa danh ngày tháng
    cell_d2_r = header_table.rows[1].cells[1]
    cell_d2_r.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for p in cell_d2_r.paragraphs:
        p._element.getparent().remove(p._element)

    ngay_ky = values.get("ngay_ky", "")
    dia_diem_ky = values.get("dia_diem_ky", "")
    try:
        parts = ngay_ky.split("/")
        if len(parts) == 3:
            ngay_str = f"ngày {parts[0]} tháng {parts[1]} năm {parts[2]}"
        else:
            ngay_str = ngay_ky
    except Exception:
        ngay_str = ngay_ky
    dia_ngay_text = f"{dia_diem_ky}, {ngay_str}" if dia_diem_ky else ngay_str

    p_dn = cell_d2_r.add_paragraph()
    _set_para_format(p_dn, WD_ALIGN_PARAGRAPH.CENTER, 6, 0, 17)
    _add_run(p_dn, dia_ngay_text, italic=True, size_pt=14)

    # ── 2.5 TÊN LOẠI VĂN BẢN (NĐ 30, Phụ lục I, Mục IV.4) ──
    # 14pt, Đậm, Hoa, căn giữa
    p_title = doc.add_paragraph()
    _set_para_format(p_title, WD_ALIGN_PARAGRAPH.CENTER, 24, 12, 17)
    _add_run(p_title, template.get("title", template["name"]).upper(), bold=True, size_pt=14)

    # ── 2.6 PREAMBLE (Căn cứ + Giới thiệu các bên) ──
    preamble_raw = template.get("preamble", "")
    preamble_rendered = render_text(preamble_raw, values)
    preamble_lines = preamble_rendered.replace("\\n", "\n").split("\n")

    for line in preamble_lines:
        line = line.strip()
        if not line:
            p = doc.add_paragraph()
            _set_para_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 3, 17)
            continue
        p = doc.add_paragraph()
        is_can_cu = line.startswith("Căn cứ") or line.startswith("căn cứ")
        is_ben = line.startswith("- BÊN") or line.startswith("-BÊN") or line.upper().startswith("BÊN")
        if is_can_cu:
            _set_para_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 6, 17, first_indent_cm=1.0)
            _add_run(p, line, italic=True, size_pt=14)
        elif is_ben:
            _set_para_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 6, 6, 17, first_indent_cm=0)
            _add_run(p, line, bold=False, size_pt=14)
        else:
            _set_para_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 6, 17, first_indent_cm=1.0)
            _add_run(p, line, size_pt=14)

    # ── 2.7 THÂN HỢP ĐỒNG (Các điều khoản) ──
    article_num = 0
    for article in template.get("body", []):
        # Kiểm tra condition
        condition = article.get("condition")
        if condition:
            cond_key = condition["key"]
            cond_val = condition["equals"]
            if values.get(cond_key) != cond_val:
                continue

        heading = article.get("heading", "").strip()
        text_raw = article.get("text", "").strip()
        text_rendered = render_text(text_raw, values)

        if heading:
            article_num += 1
            # Tiêu đề điều: "Điều X. TÊN ĐIỀU" — Đậm, 13pt, căn trái
            p_h = doc.add_paragraph()
            _set_para_format(p_h, WD_ALIGN_PARAGRAPH.LEFT, 12, 3, 17)
            heading_display = heading if heading.upper().startswith("ĐI") else f"Điều {article_num}. {heading}"
            _add_run(p_h, heading_display, bold=True, size_pt=14)

        if text_rendered:
            # Tách các dòng con trong body text
            text_lines = text_rendered.replace("\\n", "\n").split("\n")
            for tline in text_lines:
                tline = tline.strip()
                if not tline:
                    continue
                p_t = doc.add_paragraph()
                _set_para_format(p_t, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 6, 17, first_indent_cm=1.0)
                _add_run(p_t, tline, size_pt=14)

    # ── 2.8 CLOSING ──
    closing_raw = template.get("closing", "")
    if closing_raw:
        closing_rendered = render_text(closing_raw, values)
        closing_lines = closing_rendered.replace("\\n", "\n").split("\n")
        for cline in closing_lines:
            cline = cline.strip()
            if not cline:
                continue
            # Bỏ qua dòng ký tên trùng lặp để tránh chồng lấn với bảng ký tự động chuẩn NĐ 30
            sig_indicators = ["ký,", "ký và", "đóng dấu", "đại diện", "người lao động", "bên chuyển", "bên nhận", "bên tặng", "bên được", "cổ đông sáng lập", "bên a (", "bên b ("]
            if any(ind.lower() in cline.lower() for ind in sig_indicators):
                continue
            p_c = doc.add_paragraph()
            _set_para_format(p_c, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 6, 17, first_indent_cm=1.0)
            _add_run(p_c, cline, italic=True, size_pt=14)

    # ── 2.9 KHỐI KÝ (NĐ 30 — Bảng 2 cột, không viền) ──
    doc.add_paragraph()  # Khoảng cách
    p_sign_title = doc.add_paragraph()
    _set_para_format(p_sign_title, WD_ALIGN_PARAGRAPH.LEFT, 12, 6, 17)

    parties = template.get("parties", {})
    ben_a_label = parties.get("a", "Bên A")
    ben_b_label = parties.get("b", "Bên B")

    ten_nsdld = values.get("ten_nsdld", "")
    nguoi_dai_dien = values.get("nguoi_dai_dien", "")
    if " - " in nguoi_dai_dien:
        parts_nd = nguoi_dai_dien.split(" - ", 1)
        ten_nd = parts_nd[0].strip()
        chuc_vu_nd = parts_nd[1].strip().upper()
    else:
        ten_nd = nguoi_dai_dien
        chuc_vu_nd = "GIÁM ĐỐC"

    ten_nld = values.get("ten_nld", "")

    # Bảng 2 cột không viền (mỗi cột 8.0 cm để đạt 16cm chiều ngang text area)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Xóa viền
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_elem = OxmlElement(f"w:{side}")
        border_elem.set(qn("w:val"), "none")
        border_elem.set(qn("w:sz"), "0")
        border_elem.set(qn("w:space"), "0")
        border_elem.set(qn("w:color"), "auto")
        tblBorders.append(border_elem)
    tblPr.append(tblBorders)

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Cm(8.0)

    # Cột Bên A (trái)
    cell_a = table.rows[0].cells[0]
    cell_a.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    def add_cell_para(cell, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_before=0, space_after=4):
        p = cell.add_paragraph()
        _set_para_format(p, align, space_before, space_after, 17)
        if text:
            _add_run(p, text, bold=bold, italic=italic, size_pt=size)
        return p

    for p in cell_a.paragraphs:
        p._element.getparent().remove(p._element)

    add_cell_para(cell_a, f"BÊN A", bold=True, size=13)
    add_cell_para(cell_a, f"({ben_a_label})", italic=True, size=12)
    add_cell_para(cell_a, ten_nsdld.upper() if ten_nsdld else "", bold=True, size=12)
    add_cell_para(cell_a, "")
    add_cell_para(cell_a, chuc_vu_nd, bold=True, size=13)
    for _ in range(4):
        add_cell_para(cell_a, "")
    add_cell_para(cell_a, ten_nd, bold=True, size=14)

    # Cột Bên B (phải)
    cell_b = table.rows[0].cells[1]
    cell_b.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for p in cell_b.paragraphs:
        p._element.getparent().remove(p._element)

    add_cell_para(cell_b, f"BÊN B", bold=True, size=13)
    add_cell_para(cell_b, f"({ben_b_label})", italic=True, size=12)
    add_cell_para(cell_b, dia_ngay_text, italic=True, size=12)
    add_cell_para(cell_b, "")
    add_cell_para(cell_b, "")
    for _ in range(4):
        add_cell_para(cell_b, "")
    add_cell_para(cell_b, ten_nld, bold=True, size=14)

    # ── 2.10 DISCLAIMER ──
    doc.add_paragraph()
    p_disc = doc.add_paragraph()
    _set_para_format(p_disc, WD_ALIGN_PARAGRAPH.JUSTIFY, 12, 0, 16)
    _add_run(
        p_disc,
        "Lưu ý: Hợp đồng này được soạn thảo bởi legalkit-vn AI — "
        "chỉ mang tính chất tham khảo sơ bộ. Vui lòng có cán bộ pháp lý hoặc luật sư "
        "rà soát trước khi ký kết chính thức.",
        italic=True, size_pt=11, color=(128, 128, 128)
    )

    return doc


# ─────────────────────────────────────────────
# SECTION 3: Export & Main
# ─────────────────────────────────────────────

def export_docx_nd30(doc, output_path: Path):
    doc.save(output_path)
    print(f"\n[OK] Da xuat DOCX (chuan ND 30/2020): {output_path}")


def main():
    parser = argparse.ArgumentParser(description="legalkit-vn Contract Builder v2 (ND 30/2020)")
    parser.add_argument("--template", "-t", required=True, help="Ten file template JSON")
    parser.add_argument("--input", "-i", help="File JSON du lieu da dien (tuy chon)")
    parser.add_argument("--output", "-o", help="Duong dan file output .docx")
    args = parser.parse_args()

    # Load template
    try:
        template = load_template(args.template)
    except FileNotFoundError as e:
        print(f"[ERROR] {e}")
        sys.exit(1)

    # Du lieu
    if args.input:
        with open(args.input, encoding="utf-8") as f:
            values = json.load(f)
    else:
        values = collect_inputs_interactive(template)

    # Lint validation
    warnings = run_lint(template, values)
    if warnings:
        print("\n[CANH BAO PHAP LY]")
        errors = [w for w in warnings if w["severity"] == "error"]
        warns  = [w for w in warnings if w["severity"] == "warning"]
        for w in errors:
            print(f"  [ERROR] {w['message']} ({w['legalRef']})")
        for w in warns:
            print(f"  [WARNING] {w['message']} ({w['legalRef']})")
        if errors:
            print("\n[BLOCK] Co loi phap ly nghiem trong. Vui long sua truoc khi xuat.")
            sys.exit(1)

    # Render theo chuan ND 30
    print(f"\nDang render '{template['name']}' theo chuan ND 30/2020/ND-CP...")
    doc = build_docx_nd30(template, values)

    # Output path
    if args.output:
        output_path = Path(args.output)
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        output_path = OUTPUT_DIR / f"{template['id']}_{timestamp}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    export_docx_nd30(doc, output_path)

    # Summary
    lint_errors = len([w for w in warnings if w["severity"] == "error"])
    print(f"\n[TOM TAT]")
    print(f"  Template    : {template['name']}")
    print(f"  Can cu PL   : {len(template.get('legalBasis', []))} dieu")
    print(f"  Lint canh bao: {len(warnings)} ({lint_errors} loi)")
    print(f"  Chuan trinh bay: ND 30/2020/ND-CP, Phu luc I")
    print(f"  Font        : Times New Roman, Unicode")
    print(f"  Le trang    : Trai 3cm | Phai 2cm | Tren 2cm | Duoi 2cm")
    print(f"  Output      : {output_path}")


if __name__ == "__main__":
    main()
